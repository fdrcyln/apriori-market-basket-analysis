"""
Rapor Olusturma - BMM4202 Veri Madenciligi Final Odevi
"""
import os, sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUTS_DIR = os.path.join(PROJECT_ROOT, "outputs")
DATA_DIR = os.path.join(PROJECT_ROOT, "data")


def fmt(p, sz=12, bold=False, align=None, after=6):
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(sz); r.font.bold = bold

def heading(doc, txt, lv=1):
    h = doc.add_heading(txt, level=lv)
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    return h

def para(doc, txt, sz=12):
    p = doc.add_paragraph(txt); fmt(p, sz=sz); return p

def bullet(doc, txt, style="List Bullet"):
    p = doc.add_paragraph(txt, style=style)
    for r in p.runs: r.font.name="Times New Roman"; r.font.size=Pt(11)


def generate_report():
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

    # ===== KAPAK =====
    for _ in range(4): doc.add_paragraph("")
    for txt,sz,b in [("T.C.",14,True),("BALIKESİR ÜNİVERSİTESİ",16,True),("MÜHENDİSLİK FAKÜLTESİ",14,True),("BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ",14,True)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(txt); r.font.size=Pt(sz); r.font.bold=b; r.font.name="Times New Roman"
    doc.add_paragraph(""); doc.add_paragraph("")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("E-Pazaryeri Sepet Verilerinden Apriori Algoritması ile\nBirliktelik Kurallarının Çıkarılması")
    r.font.size=Pt(16); r.font.bold=True; r.font.name="Times New Roman"
    doc.add_paragraph(""); doc.add_paragraph("")
    for line in ["202213709053 Furkan DÜRCEYLAN","","BMM4202 VERİ MADENCİLİĞİ FİNAL ÖDEVİ","","Danışman: Dr. Öğr. Üyesi Kadriye ERGÜN","","BALIKESİR, HAZİRAN 2026"]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(line); r.font.size=Pt(13); r.font.name="Times New Roman"
        if "BMM4202" in line or "Danışman" in line: r.font.bold=True
    doc.add_page_break()

    # ===== 1. TANIMLAMA VE AMAC =====
    heading(doc, "1. Çalışmanın Tanımı ve Amacı")
    para(doc, "E-pazaryeri platformları günümüzde milyonlarca kullanıcıya hizmet vermekte ve her gün büyük miktarda işlem verisi üretmektedir. Bu verilerin analiz edilmesi, müşteri davranışlarının anlaşılması ve satış stratejilerinin geliştirilmesi açısından kritik öneme sahiptir. Veri madenciliği yöntemlerinden biri olan birliktelik kuralı madenciliği, bu tür büyük ölçekli işlem verilerinde ürünler arasındaki gizli ilişkileri ve birlikte satın alınma örüntülerini ortaya çıkarmak için etkili bir yaklaşım sunmaktadır.")
    para(doc, "Bu çalışmanın temel amacı, bir e-ticaret platformuna ait gerçek sepet verilerinden Apriori algoritması kullanarak birliktelik kuralları çıkarmak, bu kuralları support, confidence ve lift metrikleri üzerinden yorumlamak ve elde edilen sonuçları bir ürün öneri sistemine dönüştürmektir. Çalışma, yalnızca kural çıkarmakla sınırlı kalmayıp, çıkan kuralların anlamını ve iş dünyasındaki potansiyel kullanım alanlarını da değerlendirmektedir.")
    para(doc, "Ek olarak, web scraping yöntemiyle güncel e-pazaryeri ürün bilgileri toplanmaya çalışılmış ve bu kısım ek veri toplama denemesi olarak değerlendirilmiştir. Streamlit tabanlı interaktif bir arayüz ile kullanıcıların ürün seçerek ilişkili ürün önerilerini görmesi sağlanmıştır.")

    # ===== 2. YONTEM VE YAZILIMLAR =====
    heading(doc, "2. Kullanılan Yöntem, Program ve Yazılımlar")
    para(doc, "Bu çalışmada veri madenciliği yöntemlerinden birliktelik kuralı madenciliği (Association Rule Mining) uygulanmıştır. Çalışmada derin öğrenme teknikleri (CNN, RNN, LSTM, Transformer vb.) kesinlikle kullanılmamıştır. Bunun yerine, geleneksel veri madenciliği algoritmalarından Apriori tercih edilmiştir.")

    heading(doc, "2.1. Apriori Algoritması", lv=2)
    para(doc, "Apriori algoritması, R. Agrawal ve R. Srikant tarafından 1994 yılında önerilmiş olup büyük işlem veritabanlarında sık öğe kümelerini (frequent itemsets) keşfetmek için kullanılır. Algoritma, aşağıdan yukarıya (bottom-up) bir strateji izler: önce tekli sık öğeler bulunur, ardından bu öğeler birleştirilerek daha büyük kümeler oluşturulur. Minimum support eşiğinin altında kalan kümeler budanarak hesaplama maliyeti düşürülür.")
    para(doc, "Apriori'nin tercih edilme gerekçeleri şunlardır: (1) Algoritma kavramsal olarak anlaşılır ve yorumlanabilir sonuçlar üretir. (2) E-ticaret sepet analizi için en yaygın kullanılan yöntemlerden biridir. (3) Çıktıları doğrudan iş kurallarına dönüştürülebilir. (4) Derin öğrenme gibi kara kutu modeller yerine şeffaf ve açıklanabilir bir yapıya sahiptir.")

    heading(doc, "2.2. Değerlendirme Metrikleri", lv=2)
    para(doc, "Support (Destek): Bir kuralın veri setinde ne sıklıkla geçtiğini gösterir. Support(A→B) = P(A∪B), yani A ve B'nin birlikte geçtiği işlemlerin tüm işlemlere oranıdır. Düşük support değeri, kuralın nadir oluştuğunu ifade eder.")
    para(doc, "Confidence (Güven): A ürününü alan bir müşterinin B ürününü de alma olasılığıdır. Confidence(A→B) = P(B|A) = Support(A∪B) / Support(A). Yüksek confidence, güçlü bir yönlü ilişkiyi gösterir.")
    para(doc, "Lift (Kaldıraç): İki ürünün birlikte satın alınma eğiliminin bağımsız satın alınma olasılığına oranıdır. Lift(A→B) = Confidence(A→B) / Support(B). Lift > 1 ise pozitif ilişki (birlikte alınma eğilimi), Lift = 1 ise bağımsızlık, Lift < 1 ise negatif ilişki vardır.")

    heading(doc, "2.3. Kullanılan Teknolojiler", lv=2)
    for t in ["Python 3.x – Ana programlama dili","pandas – Veri okuma, temizleme, dönüştürme ve analiz","numpy – Sayısal hesaplamalar ve dizi işlemleri","mlxtend – Apriori algoritması (apriori, association_rules, TransactionEncoder)","matplotlib – Grafik ve görselleştirme","Streamlit – İnteraktif web arayüzü geliştirme","requests + BeautifulSoup4 – Web scraping ile veri toplama","python-docx – Bu raporun otomatik oluşturulması"]:
        bullet(doc, t)

    # ===== 3. UYGULAMA =====
    doc.add_page_break()
    heading(doc, "3. Uygulamanın Anlatılması")

    heading(doc, "3.1. Veri Setinin Tanıtılması ve Özellikleri", lv=2)
    para(doc, "Çalışmada UCI Machine Learning Repository üzerinden sağlanan ve Kaggle platformunda da erişilebilen 'Online Retail' veri seti kullanılmıştır. Bu veri seti, İngiltere merkezli bir çevrimiçi perakende şirketinin 01/12/2010 ile 09/12/2011 tarihleri arasındaki tüm işlemlerini kapsamaktadır. Şirket ağırlıklı olarak hediyelik eşya satışı yapmakta olup müşterilerinin önemli bir kısmı toptancılardan oluşmaktadır.")

    para(doc, "Veri seti toplamda 541.909 satır ve 8 sütundan oluşmaktadır. Aşağıda her sütunun adı, veri tipi ve içerdiği bilgi detaylı olarak açıklanmıştır:")

    # Veri seti ozellik tablosu
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Shading Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["Sütun Adı","Veri Tipi","Açıklama","Örnek Değer"]):
        tbl.rows[0].cells[i].text = h
    dataset_cols = [
        ("InvoiceNo","String (object)","Fatura numarası. 6 haneli. 'C' ile başlayanlar iptal edilmiş siparişleri temsil eder.","536365, C536379"),
        ("StockCode","String (object)","Ürün kodu. Her ürüne özgü 5 haneli bir tam sayı.","85123A, 71053"),
        ("Description","String (object)","Ürün adı/açıklaması. Serbest metin formatında.","WHITE HANGING HEART T-LIGHT HOLDER"),
        ("Quantity","Integer (int64)","Satın alınan ürün miktarı. Negatif değerler iade işlemlerini gösterir.","6, -2"),
        ("InvoiceDate","DateTime","İşlem tarih ve saati.","2010-12-01 08:26:00"),
        ("UnitPrice","Float (float64)","Ürün birim fiyatı (sterlin cinsinden).","2.55, 3.39"),
        ("CustomerID","Float (float64)","Müşteri kimlik numarası. Bazı kayıtlarda eksik (NaN).","17850.0, NaN"),
        ("Country","String (object)","Müşterinin bulunduğu ülke.","United Kingdom, France"),
    ]
    for name,dtype,desc,ex in dataset_cols:
        cells = tbl.add_row().cells
        cells[0].text=name; cells[1].text=dtype; cells[2].text=desc; cells[3].text=ex

    doc.add_paragraph("")
    para(doc, "Veri setinin temel istatistikleri: Toplam 541.909 işlem kaydı bulunmakta olup, bu kayıtlar 25.900 benzersiz fatura, 4.070 farklı ürün ve 4.372 benzersiz müşteriyi kapsamaktadır. Veride 38 farklı ülkeden müşteriler yer almakta olup işlemlerin büyük çoğunluğu (%91) İngiltere kaynaklıdır. Description sütununda 1.454 adet eksik (NaN) kayıt, CustomerID sütununda ise 135.080 adet eksik kayıt bulunmaktadır.")

    heading(doc, "3.2. Web Scraping ile Ek Veri Toplama Denemesi", lv=2)
    para(doc, "Projenin web scraping bileşeni olarak, scraping işlemlerine izin veren books.toscrape.com test sitesinden requests ve BeautifulSoup kütüphaneleri kullanılarak ürün bilgileri çekilmiştir. Bu işlem, ana analiz için değil; ödevde veri toplama yetkinliğini göstermek amacıyla gerçekleştirilmiştir.")
    para(doc, "Scraping sürecinde şu adımlar izlenmiştir: (1) Önce robots.txt dosyası kontrol edilmiştir. (2) HTTP GET isteği ile sayfa HTML içeriği alınmıştır. (3) BeautifulSoup ile HTML parse edilerek ürün başlıkları ve kategori bilgileri çıkarılmıştır. (4) En fazla 20 ürünle sınırlandırılmıştır. (5) Tüm süreç try-except blokları ile korunmuş, site erişim engeli durumunda otomatik olarak örnek veri oluşturulması sağlanmıştır. Çekilen veriler scraped_products.csv dosyasına product_name, category, source ve scraped_date sütunlarıyla kaydedilmiştir.")

    heading(doc, "3.3. Veri Ön İşleme Adımları", lv=2)
    para(doc, "Ham veri doğrudan Apriori algoritmasına uygun değildir. Bu nedenle aşağıdaki ön işleme adımları sırasıyla uygulanmıştır:")
    steps = [
        "Veri Yükleme: CSV dosyası pandas ile okunmuş, sütun adları otomatik algılanarak farklı veri setleriyle uyumluluk sağlanmıştır.",
        "Eksik Değer Temizleme: Description sütununda NaN olan 1.454 kayıt veri setinden çıkarılmıştır. Bu kayıtlar ürün bilgisi taşımadığından analiz için kullanılamaz.",
        "İptal İşlemlerinin Çıkarılması: InvoiceNo sütunu 'C' harfi ile başlayan 9.288 iptal edilmiş sipariş kaydı filtrelenmiştir. İptal kayıtları gerçek satın alma davranışını yansıtmaz.",
        "Negatif/Sıfır Miktar Filtreleme: Quantity ≤ 0 olan 10.624 satır çıkarılmıştır. Bu kayıtlar iade veya düzeltme işlemlerini temsil eder.",
        "Ürün Adı Standardizasyonu: Tüm ürün adları büyük harfe çevrilmiş (upper), baş/son boşluklar temizlenmiş (strip), çoklu boşluklar tek boşluğa indirgenmiştir. POSTAGE, MANUAL, BANK CHARGES gibi ürün olmayan kayıtlar çıkarılmıştır.",
        "Nadir Ürün Filtreleme: 5'ten az sipariş edilen ürünler analizden çıkarılarak gürültü azaltılmıştır.",
    ]
    for s in steps: bullet(doc, s, "List Number")
    para(doc, "Bu adımlar sonrasında orijinal 541.909 satırlık veri 528.513 satıra düşmüştür. Toplamda 13.396 satır temizlenmiştir.")

    heading(doc, "3.4. Transaction Formatına Dönüştürme", lv=2)
    para(doc, "Apriori algoritması, her satırın bir işlemi (transaction) ve her sütunun bir ürünü temsil ettiği boolean (True/False) bir matris bekler. Bu dönüşüm şu şekilde yapılmıştır: (1) Veriler InvoiceNo'ya göre gruplandırılarak her fatura için benzersiz ürün listesi oluşturulmuştur. (2) mlxtend kütüphanesinin TransactionEncoder sınıfı ile one-hot encoding uygulanmıştır. (3) Sonuçta 19.887 sipariş × 500 ürün boyutunda bir boolean matris elde edilmiştir. Performans için en sık geçen 500 ürün seçilmiştir.")

    heading(doc, "3.5. Apriori Algoritmasının Uygulanması", lv=2)
    para(doc, "mlxtend kütüphanesinin apriori() fonksiyonu ile sık öğe kümeleri bulunmuştur. Parametreler: min_support=0.02 (bir öğe kümesinin en az %2 siparişte geçmesi), use_colnames=True. Bu parametrelerle 381 adet sık öğe kümesi tespit edilmiştir. Ardından association_rules() fonksiyonu ile min_confidence=0.3 eşiğinde 153 birliktelik kuralı oluşturulmuştur.")
    para(doc, "Kodda, hiçbir kural çıkmaması durumuna karşı otomatik parametre düşürme mekanizması bulunmaktadır. min_support sırasıyla 0.01, 0.005, 0.003, 0.002, 0.001 değerlerine düşürülerek denenir. Benzer şekilde min_confidence da kademeli olarak düşürülebilir.")

    heading(doc, "3.6. Streamlit Arayüzü", lv=2)
    para(doc, "Streamlit ile geliştirilen interaktif web arayüzünde şu bileşenler yer almaktadır: (1) Veri seti istatistikleri (sipariş sayısı, ürün sayısı, kural sayısı, ortalama lift). (2) Ürün seçim kutusu – kullanıcı 500 ürün arasından seçim yapabilir. (3) Seçilen ürüne göre birliktelik kurallarına dayalı 5 ürün önerisi (confidence ve lift değerleriyle). (4) En güçlü kurallar tablosu (sıralanabilir). (5) Support-Confidence-Lift grafiği. Arayüz 'python -m streamlit run app.py' komutuyla başlatılır.")

    # ===== 4. SONUC =====
    doc.add_page_break()
    heading(doc, "4. Sonuç")
    para(doc, "Bu çalışmada 541.909 kayıtlık Online Retail veri seti üzerinde Apriori algoritması uygulanarak 153 birliktelik kuralı başarıyla çıkarılmıştır. Kuralların ortalama lift değeri 8.90, ortalama confidence değeri 0.50 olarak hesaplanmıştır. Tüm kuralların lift değeri 1'in üzerindedir, bu da bulunan ilişkilerin istatistiksel olarak anlamlı olduğunu göstermektedir.")

    # Kural tablosu ve yorumlar
    rules_path = os.path.join(OUTPUTS_DIR, "top_rules.csv")
    if os.path.exists(rules_path):
        try:
            rules_df = pd.read_csv(rules_path)
            top5 = rules_df.head(5)
            heading(doc, "4.1. En Güçlü 5 Birliktelik Kuralı", lv=2)
            tbl2 = doc.add_table(rows=1, cols=5)
            tbl2.style = "Light Shading Accent 1"
            tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i,h in enumerate(["Öncül (Antecedent)","Sonuç (Consequent)","Support","Confidence","Lift"]):
                tbl2.rows[0].cells[i].text = h
            for _,row in top5.iterrows():
                c=tbl2.add_row().cells
                c[0].text=str(row.get("antecedents_str",""))[:45]
                c[1].text=str(row.get("consequents_str",""))[:45]
                c[2].text=f"{row.get('support',0):.4f}"
                c[3].text=f"{row.get('confidence',0):.4f}"
                c[4].text=f"{row.get('lift',0):.2f}"
            doc.add_paragraph("")

            heading(doc, "4.2. Kuralların Yorumlanması", lv=2)
            # Her kural icin yorum
            for idx, row in top5.head(3).iterrows():
                ant=str(row.get("antecedents_str","X"))
                con=str(row.get("consequents_str","Y"))
                sup=row.get("support",0); conf=row.get("confidence",0); lift=row.get("lift",0)
                yorum = (
                    f"Kural: {ant} → {con}\n"
                    f"Bu kural, sepetine '{ant}' ürününü koyan müşterilerin %{conf*100:.1f} olasılıkla "
                    f"'{con}' ürününü de sepetine eklediğini göstermektedir. "
                    f"Lift değeri {lift:.2f} olup 1'den çok büyüktür; bu, iki ürün grubunun rastgele birliktelikten "
                    f"{lift:.1f} kat daha fazla birlikte satın alındığı anlamına gelir. "
                    f"Support değeri {sup:.4f} olup bu ürün çiftinin tüm siparişlerin yaklaşık %{sup*100:.1f}'inde birlikte geçtiğini ifade eder."
                )
                para(doc, yorum, sz=11)
        except Exception as e:
            para(doc, f"(Kural tablosu oluşturulamadı: {e})")

    heading(doc, "4.3. Sonuçların Genel Değerlendirmesi", lv=2)
    para(doc, "Analiz sonuçları incelendiğinde, en güçlü birliktelik kurallarının Regency serisi çay fincanları arasında yoğunlaştığı görülmektedir. Pink, Green ve Roses Regency Teacup and Saucer ürünleri birbirleriyle çok yüksek lift değerleriyle (14-18 arası) ilişkilidir. Bu durum, müşterilerin bu ürünleri set olarak satın alma eğiliminde olduğunu göstermektedir.")
    para(doc, "Pratik açıdan bu kurallar şu şekilde kullanılabilir: (1) Bir müşteri sepetine 'Pink Regency Teacup' koyduğunda, sistem otomatik olarak 'Green Regency Teacup' ve 'Roses Regency Teacup' önermelidir. (2) Bu ürünler mağazada veya web sitesinde yakın konumlandırılabilir. (3) Paket (bundle) satış kampanyaları düzenlenebilir. (4) Benzer şekilde Gardeners Kneeling Pad ve Lunch Box ürün gruplarında da çapraz satış fırsatları mevcuttur.")
    para(doc, "Yeni bir müşteri sisteme geldiğinde, sepetine eklediği ürün antecedents (öncül) kısmında aranır ve eşleşen kuralların consequents (sonuç) kısmındaki ürünler, lift ve confidence değerlerine göre sıralanarak önerilir. Örneğin, sepetinde 'DOLLY GIRL LUNCH BOX' olan bir müşteriye %63.3 confidence ve 14.0 lift ile 'SPACEBOY LUNCH BOX' önerilecektir.")

    heading(doc, "4.4. Çalışmanın Sınırlılıkları", lv=2)
    for lim in [
        "Veri seti Kaggle/UCI kaynaklı olup 2010-2011 yıllarına aittir; güncel tüketici davranışlarını tam yansıtmayabilir.",
        "Web scraping kısmı sınırlı sayıda (20) ürünle gerçekleştirilmiş olup ana analiz için kullanılmamıştır.",
        "Apriori algoritması büyük veri setlerinde hesaplama maliyeti yüksek olabilir; bu nedenle ürün sayısı 500 ile sınırlandırılmıştır.",
        "Veri setinde müşteri demografik bilgileri bulunmadığından kişiselleştirilmiş öneriler yapılamamaktadır.",
        "Support eşiği (0.02) nedeniyle nadir ama anlamlı olabilecek bazı ilişkiler gözden kaçmış olabilir.",
    ]: bullet(doc, lim)

    # ===== 5. KAYNAKLAR =====
    doc.add_page_break()
    heading(doc, "5. Kaynaklar")
    for src in [
        "[1] Daqing Chen, Online Retail Dataset, UCI Machine Learning Repository / Kaggle, https://www.kaggle.com/datasets/vijayuv/onlineretail",
        "[2] Agrawal, R., & Srikant, R. (1994). Fast Algorithms for Mining Association Rules. Proc. 20th Int. Conf. Very Large Data Bases (VLDB), 487-499.",
        "[3] Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques. 3rd Edition, Morgan Kaufmann.",
        "[4] Raschka, S. mlxtend: Providing Machine Learning Extensions. https://rasbt.github.io/mlxtend/",
        "[5] pandas Documentation. https://pandas.pydata.org/docs/",
        "[6] Streamlit Documentation. https://docs.streamlit.io/",
    ]: para(doc, src, sz=11)

    # Grafik eki
    chart_path = os.path.join(OUTPUTS_DIR, "support_confidence_lift_chart.png")
    if os.path.exists(chart_path):
        doc.add_page_break()
        heading(doc, "Ek: Support-Confidence-Lift Grafiği")
        para(doc, "Aşağıdaki grafikte en güçlü 10 birliktelik kuralının confidence ve lift değerleri karşılaştırmalı olarak gösterilmektedir.")
        doc.add_picture(chart_path, width=Inches(6))

    report_path = os.path.join(OUTPUTS_DIR, "final_report.docx")
    doc.save(report_path)
    print(f"[BASARILI] Rapor olusturuldu: {report_path}")
    return report_path

if __name__ == "__main__":
    generate_report()
